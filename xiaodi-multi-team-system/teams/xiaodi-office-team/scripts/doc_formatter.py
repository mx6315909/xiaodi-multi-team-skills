#!/usr/bin/env python3
"""
文档排版工具 - 文档管理师专用脚本
支持多种格式的文档排版：结构化 Markdown 内容、表格、列表的自动排版
利用 python-docx 生成专业排版的 docx 文件
"""

import json
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading.append(shading_elem)


def _add_run(p, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """添加带格式的 run"""
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def format_markdown_to_docx(markdown_text, output_path, options=None):
    """
    将 Markdown 文本格式化为专业排版的 docx 文档
    
    Args:
        markdown_text: Markdown 格式的文本内容
        output_path: 输出文件路径
        options: 排版选项字典，支持:
            - title: 文档标题
            - font_name: 正文字体 (默认: 宋体)
            - heading_font: 标题字体 (默认: 黑体)
            - font_size: 正文字体大小 (默认: 11)
            - heading_sizes: 各级标题大小 dict {1: 18, 2: 14, 3: 12}
            - line_spacing: 行距倍数 (默认: 1.5)
            - margins: 页边距 dict {top, bottom, left, right} in cm
            - table_style: 表格样式
            - add_toc: 是否生成目录 (默认: False)
    """
    if not HAS_DOCX:
        return {"success": False, "error": "python-docx 未安装，请运行: pip install python-docx"}
    
    opts = options or {}
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font_name = opts.get('font_name', '宋体')
    font.name = 'Times New Roman'
    font.size = Pt(opts.get('font_size', 11))
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    # 设置页边距
    margins = opts.get('margins', {'top': 2.54, 'bottom': 2.54, 'left': 3.17, 'right': 3.17})
    section = doc.sections[0]
    section.top_margin = Cm(margins.get('top', 2.54))
    section.bottom_margin = Cm(margins.get('bottom', 2.54))
    section.left_margin = Cm(margins.get('left', 3.17))
    section.right_margin = Cm(margins.get('right', 3.17))
    
    # 段落设置
    pf = style.paragraph_format
    pf.line_spacing = opts.get('line_spacing', 1.5)
    pf.space_after = Pt(4)
    
    heading_sizes = opts.get('heading_sizes', {1: 18, 2: 14, 3: 12, 4: 11})
    heading_font = opts.get('heading_font', '黑体')
    
    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_headers = []
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 空行
        if not stripped:
            if in_table and table_rows:
                # 输出表格
                if table_headers:
                    t = doc.add_table(rows=1+len(table_rows), cols=len(table_headers))
                    t.style = opts.get('table_style', 'Light Grid Accent 1')
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ci, h in enumerate(table_headers):
                        cell = t.rows[0].cells[ci]
                        cell.text = h
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.bold = True
                    for ri, row in enumerate(table_rows):
                        for ci in range(min(len(row), len(table_headers))):
                            t.rows[ri+1].cells[ci].text = str(row[ci])
                doc.add_paragraph()
                in_table = False
                table_headers = []
                table_rows = []
            else:
                doc.add_paragraph()
            i += 1
            continue
        
        # 标题
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            if 1 <= level <= 4:
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if level == 1 else RGBColor(0x33, 0x33, 0x33)
                    run.font.name = heading_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), heading_font)
                    run.font.size = Pt(heading_sizes.get(level, 12))
            else:
                p = doc.add_paragraph(text)
            i += 1
            continue
        
        # 分隔线
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, '─' * 40, color=(0xBB, 0xBB, 0xBB))
            i += 1
            continue
        
        # 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # 分隔行
            if re.match(r'^[\|\s\-\:]+$', stripped):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        
        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:].strip()
            _render_bold_text(p, text)
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # 有序列表
        if re.match(r'^\d+[\.\、）\)]\s', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            m = re.match(r'^(\d+[\.\、）\)]\s*\*?\*?)(.*)', stripped)
            if m:
                _add_run(p, m.group(1), bold=True)
                _render_bold_text(p, m.group(2))
            else:
                p.add_run(stripped)
            i += 1
            continue
        
        # 引用
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            text = stripped.lstrip('>').strip()
            _add_run(p, text, italic=True, color=(0x66, 0x66, 0x66))
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # 普通段落
        if stripped:
            p = doc.add_paragraph()
            _render_bold_text(p, stripped)
        i += 1
    
    # 清理最后未输出的表格
    if in_table and table_rows and table_headers:
        t = doc.add_table(rows=1+len(table_rows), cols=len(table_headers))
        t.style = opts.get('table_style', 'Light Grid Accent 1')
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(table_headers):
            cell = t.rows[0].cells[ci]
            cell.text = h
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
        for ri, row in enumerate(table_rows):
            for ci in range(min(len(row), len(table_headers))):
                t.rows[ri+1].cells[ci].text = str(row[ci])
    
    doc.save(output_path)
    return {"success": True, "output": str(output_path)}


def _render_bold_text(paragraph, text):
    """渲染带 **加粗** 的文本"""
    if '**' in text:
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                _add_run(paragraph, part[2:-2], bold=True)
            else:
                paragraph.add_run(part)
    else:
        paragraph.add_run(text)


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='文档排版工具')
    parser.add_argument('input', nargs='?', help='输入文件路径（如省略则从 stdin 读取）')
    parser.add_argument('-o', '--output', default='output.docx', help='输出文件路径')
    parser.add_argument('--title', '-t', help='文档标题')
    parser.add_argument('--json', '-j', action='store_true', help='JSON 模式输出')
    parser.add_argument('--font-size', type=int, default=11, help='正文字体大小')
    parser.add_argument('--font-name', default='宋体', help='正文字体')
    parser.add_argument('--heading-font', default='黑体', help='标题字体')
    parser.add_argument('--line-spacing', type=float, default=1.5, help='行距')
    parser.add_argument('--no-bullets', action='store_true', help='不使用列表符号（用段落替代）')
    parser.add_argument('--style', default='default', choices=['default', 'compact', 'report'], help='排版风格')
    
    args = parser.parse_args()
    
    # 读取输入
    if args.input:
        input_path = Path(args.input)
        if not input_path.exists():
            print(json.dumps({"success": False, "error": f"文件不存在: {args.input}"}))
            sys.exit(1)
        text = input_path.read_text(encoding='utf-8')
    else:
        text = sys.stdin.read()
    
    if not text.strip():
        print(json.dumps({"success": False, "error": "输入内容为空"}))
        sys.exit(1)
    
    # 构建选项
    options = {
        'title': args.title,
        'font_name': args.font_name,
        'heading_font': args.heading_font,
        'font_size': args.font_size,
        'line_spacing': args.line_spacing,
    }
    
    # 根据风格调整
    if args.style == 'compact':
        options['line_spacing'] = 1.15
        options['font_size'] = 10
        options['margins'] = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
    elif args.style == 'report':
        options['line_spacing'] = 1.8
        options['heading_sizes'] = {1: 22, 2: 16, 3: 14, 4: 12}
        options['margins'] = {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6}
    
    result = format_markdown_to_docx(text, args.output, options)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
